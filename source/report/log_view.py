import re
import copy
from graphviz import Source
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from pydocx import PyDocX
# import mammoth
from os.path import basename
from subprocess import call
import os
import xlwt


# def set_style(name, height, bold=False, format_str='', align='center'):
#     style = xlwt.XFStyle()  # 初始化样式
#
#     font = xlwt.Font()  # 为样式创建字体
#     font.name = name  # 'Times New Roman'
#     font.bold = bold
#     font.height = height
#
#     borders = xlwt.Borders()  # 为样式创建边框
#     borders.left = 2
#     borders.right = 2
#     borders.top = 0
#     borders.bottom = 2
#
#     alignment = xlwt.Alignment()  # 设置排列
#     if align == 'center':
#         alignment.horz = xlwt.Alignment.HORZ_CENTER
#         alignment.vert = xlwt.Alignment.VERT_CENTER
#     else:
#         alignment.horz = xlwt.Alignment.HORZ_LEFT
#         alignment.vert = xlwt.Alignment.VERT_BOTTOM
#
#     style.font = font
#     style.borders = borders
#     style.num_format_str = format_str
#     style.alignment = alignment
#
#     return style
#
# class Generate_excel(object):
#     wbk = ""
#     datasheet = ""
#     head_title = ""
#     got_log = ""
#     def __init__(self, got_log_path, heap_log_path):
#         self.wbk = xlwt.Workbook(encoding="utf-8")
#         self.datasheet = self.wbk.add_sheet("sheet1",cell_overwrite_ok=True)
#         self.datasheet.write_merge(0,3,0,9,'EXPORT REPORT',set_style("Times New Roman",640,True))
#         # self.wbk.save('test.xls')
#         self.got_log = View_got_log(got_log_path).get_got_change()
#
#     def save_excel(self):
#         self.wbk.save('test.xls')
#
#     def add_got_analy(self):
#         rowNum = self.datasheet.last_used_row
#         self.datasheet.write_merge(rowNum+1,rowNum+3,0,9,"got analysis",set_style("Times New Roman",480,True,align="Left"))
#         self.datasheet.write_merge(rowNum+4,rowNum+4,0,9,"")
#         rowNum += 5
#         self.datasheet.write(rowNum, 0, "got mismatch")
#         self.datasheet.write(rowNum, 1, "which function")
#         rowNum += 1
#         for val1, val2 in (self.got_log):
#             self.datasheet.write(rowNum, 0, val1)
#             self.datasheet.write(rowNum, 1, val2)
#             rowNum += 1

class Generate_doc(object):
    document = ""
    got_log = ""
    def __init__(self, got_log_path, heap_log_path):
        self.document = Document()
        head = self.document.add_heading('\t\t\tEXPLOIT REPORT', 0)
        self.got_log = View_got_log(got_log_path).get_got_change()
        View_heap_log(heap_log_path).gen_heap_change_png()

    def docx2pdf(self, docx_path):
        '''
        tranfser docx to pdf
        :param docx_path:
        :return:
        '''
        if docx_path.endswith('.docx'):
            order = 'libreoffice --invisible --convert-to pdf %s 1>/dev/null 2>&1' % docx_path
        else:
            print('Error, file type does not match!')
            return
        call(order, shell=True)

    def save_report(self):
        '''
        save file
        :return:
        '''
        self.document.save('./report.docx')
        self.docx2pdf("./report.docx")
        html = PyDocX.to_html("./report.docx")
        f = open("./report.html", 'w', encoding="utf-8")
        f.write(html)
        f.close()


    def heap_analy_report(self):
        '''
        add heap analysis content
        :return:
        '''
        self.document.add_heading('heap analysis', 1)
        self.document.add_heading('1.got analysis', 2)
        self.document.add_paragraph("got change:")
        got_table = self.document.add_table(rows=0, cols=2)
        cells = got_table.add_row().cells
        cells[0].text = "got mismatch"
        cells[1].text = "which function"
        for i in range(len(self.got_log)):
            cells = got_table.add_row().cells
            cells[0].text = self.got_log[i][0]
            cells[1].text = self.got_log[i][1]

        self.document.add_heading("2.heap analysis", 2)
        self.document.add_paragraph("heap change graph")
        pic = self.document.add_picture("./HeapChange.png", width=Inches(6.0))
        # pic.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER.


class view_leak_log(object):
    def __init__(self, path):
        self.path = path

    def get_leak_table(self):
        '''
        :return: a list whose format is for jinja2
        '''
        def parse_info_from_log(path):
            '''
            parse the leak analysis log file
            :param path: leak analysis log file path
            :return: a list whose format is for jinja2
            '''

            leak_list = []
            f = open(path, "r")
            while True:
                node_info = {}
                line = f.readline()
                if not line:
                    break
                obj = re.match("Found leaked addr: (.*) (.*) in lib (.*)", line)
                if obj:
                    node_info['name'] = obj.group(1)
                    node_info['addr'] = obj.group(2)
                    node_info['lib'] = obj.group(3)
                    leak_list.append(copy.deepcopy(node_info))
            f.close()
            return leak_list

        return parse_info_from_log(self.path)

class view_got_log(object):
    def __init__(self, path):
        self.path = path

    def get_got_change(self):
        '''
        :return: a list whose format is for jinja2
        '''

        def parse_info_from_log(path):
            '''
            parse the got analysis log file
            :param path: got analysis log file path
            :return: a list whose format is for jinja2
            '''
            f = open(path, "r")
            got_change = []
            lines = f.readlines()
            for i in range(len(lines)):
                node_info = {}
                misobj = re.match("Found got mismatch: (.*)", lines[i])
                if misobj:
                    node_info['mismatch'] = misobj.group(1)
                    if i == len(lines) - 1:
                        continue
                    if lines[i + 1].startswith("which is func"):
                        obj = re.match("which is func (.*)", lines[i + 1])
                        node_info['function'] = obj.group(1)
                    else:
                        # node_info.append("")
                        node_info['function'] = ""
                    got_change.append(copy.deepcopy(node_info))
            f.close()
            return got_change

        return parse_info_from_log(self.path)


def overflow_heap_info(node_info, overflow_info):
    '''
    judge which heap is overflow point
    :param node_info: a list of heap info
    :param overflow_info: a node which is overflow
    :return: the node which is overflow point
    '''
    for info in node_info:
        start = int(info[0], 16)-0x10
        end = start + int(info[1], 16)
        overflow_start = int(overflow_info[0], 16)
        if overflow_start >= start and overflow_start <= end:
            info[2] = overflow_info[2]
            break
    return node_info

def free_heap_info(node_info, free_info):
    '''
    delete the released heap
    :param node_info: a list of heap info
    :param free_info: a node which will be free
    :return: the node which will be free in heap info
    '''
    success = False
    index = 0
    for info in node_info:
        if info[0] == free_info[0]:
            success = True
            break
        index += 1
    if success:
        del node_info[index]
        return node_info
    else:
        return node_info

class view_heap_log(object):

    def __init__(self, path):
        self.path = path


    def gen_heap_change_png(self):
        '''
        generate dot file and png file
        :return:
        '''

        def parse_heap_change_log(path):
            '''
            Record heap change information
            :param path: heap analysis log file path
            :return: a list of heap info change
            '''

            heap_infos = []
            f = open(path, "r")
            node_info = []
            while True:
                line = f.readline()
                if not line:
                    break
                mallocobj = re.match("Malloc called with size (.*), returns addr (.*)", line)
                if mallocobj:
                    # print("malloc:" + mallocobj.group(1) + " " + mallocobj.group(2))
                    malloc_info = [mallocobj.group(2), mallocobj.group(1), mallocobj.group()]
                    node_info.append(malloc_info)
                    heap_infos.append([mallocobj.group(), copy.deepcopy(node_info)])
                    continue

                callocobj = re.match("Calloc called with size (.*), returns addr (.*)", line)
                if callocobj:
                    # print("calloc:" + callocobj.group(1) + " " + callocobj.group(2))
                    calloc_info = [callocobj.group(2), callocobj.group(1), callocobj.group()]
                    node_info.append(calloc_info)
                    heap_infos.append([callocobj.group(), copy.deepcopy(node_info)])
                    continue

                freeobj = re.match("Free called to free (.*) with size (.*)", line)
                if freeobj:
                    # print("free:" + freeobj.group(1) + " " + freeobj.group(2))
                    free_info = [freeobj.group(1), freeobj.group(2), freeobj.group()]
                    node_info = free_heap_info(node_info, free_info)
                    heap_infos.append([freeobj.group(), copy.deepcopy(node_info)])
                    continue

                overflowobj = re.match(
                    "Found overflow in chunk at 0x6031b0, size 0x20 with write starts at <BV64 (.*)>, size <BV64 (.*)>!\n",
                    line)
                if overflowobj:
                    # print("overflow:" + overflowobj.group(1) + " " + overflowobj.group(2))
                    overlength = f.readline()
                    overflow_info = [overflowobj.group(1), overflowobj.group(2), overflowobj.group() + overlength]
                    node_info = overflow_heap_info(node_info, overflow_info)
                    heap_infos.append([overflowobj.group() + overlength, copy.deepcopy(node_info)])
            f.close()
            # print(self.heap_infos)
            return heap_infos


        heap_infos = parse_heap_change_log(self.path)
        head_dot = '''digraph G {n0[shape=reocord,label="......"]'''
        tail_dot = "}"
        label_dot = ""
        edge_dot = ""
        index = 0
        for heap_info in heap_infos:
            index += 1
            if len(heap_info[1]) == 0:
                label_dot += '''n%s[shape=record,label="......"]''' % (index)
                edge_dot += '''n%s->n%s[label="%s"]''' % (index-1, index, heap_info[0])
                continue
            content = '''n%s[shape=none, label=<<table border="0" cellborder="1" cellspacing="0" cellpadding="4">''' % (index)
            is_overflow = False
            for info in heap_info[1]:
                if info[2].startswith("Found overflow"):
                    content += '''<tr><td bgcolor="lightgrey"><font color="red">%s size:%s</font></td></tr>''' % (info[0], info[1])
                    is_overflow = True
                    continue
                content += '''<tr><td>%s size:%s</td></tr>''' % (info[0], info[1])
            content += '''</table>>]'''
            label_dot += content
            edge_dot += '''n%s->n%s[label="%s"]''' % (index-1, index, heap_info[0])

        dot = head_dot + label_dot + edge_dot + tail_dot
        t = Source(dot)
        t.save("HeapChange.dot")
        os.system("dot ./HeapChange.dot -T png -o ./HeapChange.png")
        return os.path.join(os.getcwd(), 'HeapChange.png')